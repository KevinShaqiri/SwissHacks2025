import os
import json
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

# === Step 1: Load API Key from .env ===
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# === Step 2: Extract text from .docx file ===
def docx_to_text(docx_path: str):
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text.strip())

    return "\n".join(lines)


# === Step 3: Send to OpenAI and parse JSON ===
def extract_json_from_docx_text(docx_text: str):
    system_prompt = """
You are a data parser. The user provides raw text from a structured onboarding form (.docx) with fields like personal info, financials, and checkboxes.

Extract this information into a clean JSON using specific field names. If a field is missing, return an empty string or empty list.

Only return the JSON. Do not explain anything. Do not wrap it in markdown or triple quotes.
"""

    user_prompt = f"""Here is the full extracted text from the .docx file:

{docx_text}

Now extract and return the JSON with the following format and field names:

{{
  "last_name": "...",
  "first_middle_name": "...",
  "address": "...",
  "country_of_domicile": "...",
  "date_of_birth": "...",
  "nationality": "...",
  "passport_no_or_unique_id": "...",
  "id_type": "...",
  "id_issue_date": "...",
  "id_expiry_date": "...",
  "gender": "...",
  "communication_medium_telephone": "...",
  "communication_medium_email": "...",
  "politically_exposed": "...",
  "marital_status": "...",
  "highest_education_attained": "...",
  "education_history": "...",
  "current_employment_function": "...",
  "employment_since": "...",
  "employer_name": "...",
  "position": "...",
  "self_employed": "...",
  "self_employed_since": "...",
  "self_employed_company_name": "...",
  "self_employed_ownership": "...",
  "currently_not_employed_since": "...",
  "retired_since": "...",
  "homemaker_since": "...",
  "student_since": "...",
  "student_country": "...",
  "diplomat_since": "...",
  "diplomat_country": "...",
  "military_since": "...",
  "military_country": "...",
  "other_profession_since": "...",
  "unknown_profession": "...",
  "industry": "...",
  "industry_description": "...",
  "main_business_activities": "...",
  "countries_of_business": "...",
  "total_wealth_estimated": "...",
  "origin_of_wealth": ["...", "..."],
  "origin_of_wealth_description": "...",
  "source_of_information": "...",
  "estimated_assets_real_estate": "...",
  "estimated_assets_business": "...",
  "estimated_assets_investments": "...",
  "estimated_assets_deposits": "...",
  "estimated_assets_equity": "...",
  "estimated_assets_fixed_income": "...",
  "estimated_assets_structured_products": "...",
  "estimated_assets_alternative_investments": "...",
  "estimated_assets_investment_funds": "...",
  "estimated_assets_insurance": "...",
  "estimated_assets_other": "...",
  "estimated_total_income": "...",
  "country_of_main_income": "...",
  "account_number": "...",
  "commercial_account": "...",
  "investment_risk_profile": "...",
  "type_of_mandate": "...",
  "investment_experience": "...",
  "investment_horizon": "...",
  "expected_transactional_behavior": "...",
  "preferred_markets": "...",
  "total_asset_under_management": "...",
  "aum_to_transfer_to_bjb": "..."
}}"""

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt.strip()},
            {"role": "user", "content": user_prompt.strip()}
        ],
        temperature=0
    )

    content = response.choices[0].message.content.strip()

    try:
        return json.loads(content)
    except json.JSONDecodeError:
        print("⚠️ Could not parse JSON. Model returned:")
        print(content)
        return {}


# === Step 4: Main script ===
if __name__ == "__main__":
    docx_path = "./data/profile.docx"
    output_path = "./data/profile_docx.json"

    text = docx_to_text(docx_path)
    json_result = extract_json_from_docx_text(text)

    print(f"Saving JSON to {output_path}")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(json_result, f, indent=2, ensure_ascii=False)
